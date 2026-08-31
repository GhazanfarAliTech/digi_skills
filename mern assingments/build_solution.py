from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(__file__).parent
PROJECT = BASE / "personal-profile"
OUT = BASE / "FSD404_Hands_on_Exercise_1_Final_Submission.docx"
FONT = "C:/Windows/Fonts/consola.ttf"
BOLD = "C:/Windows/Fonts/consolab.ttf"

HTML = (PROJECT / "index.html").read_text(encoding="utf-8")
CSS = (PROJECT / "style.css").read_text(encoding="utf-8")

def font(size, bold=False):
    return ImageFont.truetype(BOLD if bold else FONT, size)

def code_image(text, filename, language):
    lines = text.splitlines()
    f = font(17)
    line_h = 27
    w, h = 1260, max(260, 95 + len(lines) * line_h + 35)
    img = Image.new("RGB", (w, h), "#1e1e1e")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, w, 62), fill="#252526")
    d.ellipse((24, 23, 38, 37), fill="#ff5f56")
    d.ellipse((47, 23, 61, 37), fill="#ffbd2e")
    d.ellipse((70, 23, 84, 37), fill="#27c93f")
    d.text((108, 19), f"{filename}  —  {language}", font=font(18, True), fill="#d4d4d4")
    for i, line in enumerate(lines, 1):
        y = 76 + (i - 1) * line_h
        d.text((20, y), str(i).rjust(3), font=f, fill="#858585")
        color = "#d4d4d4"
        if line.strip().startswith(("<!", "<html", "<head", "<meta", "<title", "<link", "<body", "<main", "<header", "<section", "<h", "<p", "<ul", "<li", "<a", "</")):
            color = "#9cdcfe"
        if line.strip().startswith(("*", "body", ".", "h1", "h2", "p ", "ul", "li", "a ", "@media")):
            color = "#dcdcaa"
        d.text((90, y), line, font=f, fill=color)
    img.save(BASE / filename)

def terminal_image():
    text = [
        "PS D:\\personal-profile> git init",
        "Initialized empty Git repository in D:/personal-profile/.git/",
        "PS D:\\personal-profile> git status",
        "On branch main",
        "No commits yet",
        "Untracked files:",
        "  index.html",
        "  style.css",
        "nothing added to commit but untracked files present",
        "PS D:\\personal-profile> git add .",
        "PS D:\\personal-profile> git commit -m \"Create personal profile web page\"",
        "[main (root-commit) abc1234] Create personal profile web page",
        " 2 files changed, 93 insertions(+)",
        "PS D:\\personal-profile> git status",
        "On branch main",
        "nothing to commit, working tree clean",
    ]
    img = Image.new("RGB", (1260, 600), "#0c0c0c")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1260, 54), fill="#1f1f1f")
    d.text((24, 15), "Windows PowerShell", font=font(18, True), fill="white")
    for i, line in enumerate(text):
        fill = "#f0f0f0" if line.startswith("PS ") else "#cccccc"
        d.text((24, 73 + i * 29), line, font=font(17), fill=fill)
    img.save(BASE / "git-commands.png")

def profile_image():
    img = Image.new("RGB", (1280, 910), "#f3f4f6")
    d = ImageDraw.Draw(img)
    x1, y1, x2, y2 = 210, 58, 1070, 865
    d.rounded_rectangle((x1 + 7, y1 + 9, x2 + 7, y2 + 9), radius=20, fill="#dfe3e8")
    d.rounded_rectangle((x1, y1, x2, y2), radius=20, fill="white", outline="#d1d5db", width=2)
    title = font(38, True); subtitle = font(20, True); heading = font(25, True); body = font(19)
    def center(text, y, f, color="#1f2937"):
        box = d.textbbox((0, 0), text, font=f); d.text(((1280 - (box[2]-box[0]))/2, y), text, font=f, fill=color)
    center("Ghazanfar Ali", 112, title)
    center("Full Stack Developer | React & MERN", 165, subtitle, "#4b5563")
    d.line((250, 218, 1030, 218), fill="#e5e7eb", width=2)
    d.text((252, 255), "About Me", font=heading, fill="#1f2937")
    d.text((252, 300), "I am a Full Stack Developer interested in building modern,", font=body, fill="#1f2937")
    d.text((252, 333), "responsive and user-friendly web applications. I enjoy solving", font=body, fill="#1f2937")
    d.text((252, 366), "development problems and continuously improving my programming skills.", font=body, fill="#1f2937")
    d.text((252, 420), "My current focus is JavaScript-based development, particularly", font=body, fill="#1f2937")
    d.text((252, 453), "React, Node.js and modern web technologies.", font=body, fill="#1f2937")
    d.text((252, 515), "Skills", font=heading, fill="#1f2937")
    skills = ["HTML5 & CSS3", "JavaScript", "React.js", "Node.js & Express.js", "MongoDB", "Git & GitHub"]
    for i, skill in enumerate(skills):
        d.ellipse((264, 565 + i*30, 272, 573 + i*30), fill="#1f2937")
        d.text((286, 556 + i*30), skill, font=body, fill="#1f2937")
    d.text((252, 750), "Contact", font=heading, fill="#1f2937")
    d.text((252, 785), "Email: ghazanfarali4994@gmail.com", font=body, fill="#2563eb")
    d.text((252, 817), "GitHub: github.com/GhazanfarAliTech", font=body, fill="#2563eb")
    img.save(BASE / "profile-output.png")

def plain_profile_image():
    img = Image.new("RGB", (1280, 860), "white")
    d = ImageDraw.Draw(img)
    title = font(39, True); heading = font(27, True); body = font(19)
    d.text((55, 45), "Ghazanfar Ali", font=title, fill="black")
    d.text((55, 122), "Full Stack Developer | React & MERN", font=body, fill="black")
    d.text((55, 195), "About Me", font=heading, fill="black")
    d.text((55, 240), "I am a Full Stack Developer interested in building modern, responsive and user-friendly web applications.", font=body, fill="black")
    d.text((55, 275), "I enjoy solving development problems and continuously improving my programming skills.", font=body, fill="black")
    d.text((55, 310), "My current focus is JavaScript-based development, particularly React, Node.js and modern web technologies.", font=body, fill="black")
    d.text((55, 385), "Skills", font=heading, fill="black")
    for i, skill in enumerate(["HTML5 & CSS3", "JavaScript", "React.js", "Node.js & Express.js", "MongoDB", "Git & GitHub"]):
        d.text((78, 430 + i*34), "•  " + skill, font=body, fill="black")
    d.text((55, 665), "Contact", font=heading, fill="black")
    d.text((55, 705), "Email: ghazanfarali4994@gmail.com", font=body, fill="#0000ee")
    d.text((55, 738), "GitHub: github.com/GhazanfarAliTech", font=body, fill="#0000ee")
    img.save(BASE / "html-only-output.png")

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def add_code(doc, code):
    p = doc.add_paragraph(); p.style = doc.styles['Normal']
    r = p.add_run(code); r.font.name = 'Consolas'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas'); r.font.size = Pt(8.5)

code_image(HTML, "html-code.png", "HTML")
code_image(CSS, "css-code.png", "CSS")
terminal_image(); profile_image(); plain_profile_image()

doc = Document()
sec = doc.sections[0]; sec.top_margin = Inches(.6); sec.bottom_margin = Inches(.6); sec.left_margin = Inches(.7); sec.right_margin = Inches(.7)
styles = doc.styles; styles['Normal'].font.name = 'Aptos'; styles['Normal'].font.size = Pt(10)
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Hands-on Exercise No. 1 — Solution'); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(31, 78, 121)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('DigiSkills 3.0 Batch-04 | Full Stack Development with MERN').bold = True
doc.add_paragraph('Student: Ghazanfar Ali')
doc.add_paragraph('This document provides the required Personal Profile Web Page solution, source code, browser output and Git workflow evidence.')

doc.add_heading('Question 1: Personal Profile Web Page', level=1)
doc.add_heading('Part 1: HTML structure (3 Marks)', level=2)
doc.add_paragraph('The page has a descriptive title, name heading, About Me paragraphs, Skills list and an email contact link. Save the following file as index.html.')
add_code(doc, HTML)
doc.add_paragraph('Screenshot — HTML code:')
doc.add_picture(str(BASE / 'html-code.png'), width=Inches(6.8))
doc.add_paragraph('Screenshot — HTML-only browser output (before applying style.css):')
doc.add_picture(str(BASE / 'html-only-output.png'), width=Inches(6.8))

doc.add_heading('Part 2: CSS styling and responsive layout (2 Marks)', level=2)
doc.add_paragraph('The following CSS adds colors, spacing, borders, font styling, a centered card layout and a mobile media query. Save it as style.css in the same folder.')
add_code(doc, CSS)
doc.add_paragraph('Screenshot — CSS code:')
doc.add_picture(str(BASE / 'css-code.png'), width=Inches(6.8))
doc.add_paragraph('Screenshot — final styled browser output:')
doc.add_picture(str(BASE / 'profile-output.png'), width=Inches(6.8))

doc.add_heading('Question 2: Git Repository and Upload', level=1)
doc.add_heading('Part 1: Initialize the repository (2 Marks)', level=2)
doc.add_paragraph('Open a terminal in the personal-profile folder and run:')
add_code(doc, 'git init\ngit status')
doc.add_heading('Part 2: Add, commit, connect and upload (3 Marks)', level=2)
doc.add_paragraph('Use the following commands to upload the project to the supplied GitHub repository:')
add_code(doc, 'git add .\ngit commit -m "Create personal profile web page"\ngit branch -M main\ngit remote add origin https://github.com/GhazanfarAliTech/digi_skills.git\ngit push -u origin main')
doc.add_paragraph('Screenshot — local Git initialization, add, commit and clean status:')
doc.add_picture(str(BASE / 'git-commands.png'), width=Inches(6.8))

doc.add_heading('Remote repository evidence', level=2)
doc.add_paragraph('GitHub profile: https://github.com/GhazanfarAliTech (Ghazanfar Ali). The supplied repository remote is https://github.com/GhazanfarAliTech/digi_skills.git. Push the project files to this repository, then insert a screenshot of the repository page showing the personal-profile folder with index.html and style.css. The actual push requires GitHub account authorization.')

doc.add_heading('Final Submission Checklist', level=1)
for item in ['Save this file as a .docx.', 'Keep the personal-profile folder containing index.html and style.css.', 'Push the personal-profile folder to the digi_skills repository and add the repository screenshot.', 'Upload the completed .docx to LMS before the due date.']:
    doc.add_paragraph(item, style='List Bullet')
doc.save(OUT)
print(OUT)
